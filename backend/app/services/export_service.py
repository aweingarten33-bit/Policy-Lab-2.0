"""
Export Service v3 — Source-Grounded Compliance Intelligence System.

Generates professional .docx reports with full source attribution and
verification status for both single gap analysis and the Complete
Compliance Action Package (all 7 outputs).

Key enhancements over v2:
  - Source attribution labels on every finding, change, task, and checklist item
  - Verification status indicators (verified, partially verified, unverified, contradicted)
  - KB Sources Used section on the cover page
  - Live Research Used indicator
  - Overall Verification Summary section
  - 4-tier source type labels (Verified Source, Retrieved Source, Live Research, Model Inference)
  - Source Attribution Legend explaining the labeling system
"""

import io
import logging
from datetime import datetime
from typing import Optional, List

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from app.models.schemas import (
    AnalysisResult, GapRow, RiskLevel, ExportFormat,
    ComplianceActionPackage, RewrittenPolicy, RedlineChange,
    RemediationPlan, RemediationPhase, RemediationTask,
    BoardSummary, ImplementationChecklist, ChecklistItem,
    SourceAttribution, SourceType, VerificationStatus,
)

logger = logging.getLogger(__name__)

# ── Color constants ──
COLOR_DARK_NAVY = RGBColor(0x1A, 0x1A, 0x2E)
COLOR_ACCENT_ORANGE = RGBColor(0xF5, 0x9E, 0x0B)
COLOR_CRITICAL = RGBColor(0xC0, 0x39, 0x2B)
COLOR_HIGH = RGBColor(0xE6, 0x7E, 0x22)
COLOR_MODERATE = RGBColor(0xF3, 0x9C, 0x12)
COLOR_LOW = RGBColor(0x29, 0x80, 0xB9)
COLOR_COMPLIANT = RGBColor(0x27, 0xAE, 0x60)
COLOR_GRAY = RGBColor(0x66, 0x66, 0x66)
COLOR_LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_BLACK = RGBColor(0x1A, 0x1A, 0x1A)

# ── Source Attribution Colors ──
COLOR_VERIFIED_SOURCE = RGBColor(0x27, 0xAE, 0x60)      # Green — cross-checked against source
COLOR_RETRIEVED_SOURCE = RGBColor(0x29, 0x80, 0xB9)     # Blue — from curated knowledge base
COLOR_LIVE_RESEARCH = RGBColor(0x8E, 0x44, 0xAD)        # Purple — from controlled web research
COLOR_MODEL_INFERENCE = RGBColor(0xE6, 0x7E, 0x22)      # Orange — LLM training data, NOT verified

# ── Verification Status Colors ──
COLOR_STATUS_VERIFIED = RGBColor(0x27, 0xAE, 0x60)       # Green
COLOR_STATUS_PARTIAL = RGBColor(0xF3, 0x9C, 0x12)        # Amber
COLOR_STATUS_UNVERIFIED = RGBColor(0xE6, 0x7E, 0x22)     # Orange
COLOR_STATUS_CONTRADICTED = RGBColor(0xC0, 0x39, 0x2B)   # Red

RISK_COLORS = {
    RiskLevel.critical: COLOR_CRITICAL,
    RiskLevel.high: COLOR_HIGH,
    RiskLevel.moderate: COLOR_MODERATE,
    RiskLevel.low: COLOR_LOW,
    RiskLevel.compliant: COLOR_COMPLIANT,
}

RISK_LABELS = {
    RiskLevel.critical: "CRITICAL",
    RiskLevel.high: "HIGH",
    RiskLevel.moderate: "MODERATE",
    RiskLevel.low: "LOW",
    RiskLevel.compliant: "COMPLIANT",
}

FONT_FAMILY = "Times New Roman"

# ── Source type display mappings ──

SOURCE_TYPE_LABELS = {
    SourceType.verified_source: "Verified Source",
    SourceType.retrieved_source: "Retrieved Source",
    SourceType.live_research: "Live Research",
    SourceType.model_knowledge: "Model Inference",
}

SOURCE_TYPE_COLORS = {
    SourceType.verified_source: COLOR_VERIFIED_SOURCE,
    SourceType.retrieved_source: COLOR_RETRIEVED_SOURCE,
    SourceType.live_research: COLOR_LIVE_RESEARCH,
    SourceType.model_knowledge: COLOR_MODEL_INFERENCE,
}

SOURCE_TYPE_BG_HEX = {
    SourceType.verified_source: "E8F5E9",    # Light green
    SourceType.retrieved_source: "E3F2FD",    # Light blue
    SourceType.live_research: "F3E5F5",       # Light purple
    SourceType.model_knowledge: "FFF8E1",     # Light amber
}

VERIFICATION_STATUS_LABELS = {
    VerificationStatus.verified: "Verified",
    VerificationStatus.partially_verified: "Partially Verified",
    VerificationStatus.unverified: "Unverified",
    VerificationStatus.contradicted: "Contradicted",
}

VERIFICATION_STATUS_COLORS = {
    VerificationStatus.verified: COLOR_STATUS_VERIFIED,
    VerificationStatus.partially_verified: COLOR_STATUS_PARTIAL,
    VerificationStatus.unverified: COLOR_STATUS_UNVERIFIED,
    VerificationStatus.contradicted: COLOR_STATUS_CONTRADICTED,
}


def _set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_styled_paragraph(doc, text: str, bold: bool = False, size: int = 11,
                           color: RGBColor = COLOR_BLACK, alignment=None,
                           space_after: int = 120, space_before: int = 0,
                           italic: bool = False, font_family: str = FONT_FAMILY):
    """Add a paragraph with consistent styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font_family
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after / 20)
    pf.space_before = Pt(space_before / 20)
    return p


def _add_horizontal_rule(doc, color: str = "F59E0B"):
    """Add a styled horizontal rule."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="{color}"/></w:pBdr>')
    pPr.append(pBdr)


def _add_disclaimer_box(doc):
    """Add the AI disclaimer box with source attribution awareness."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_shading(cell, "FFF3F3")

    p = cell.paragraphs[0]
    run = p.add_run("IMPORTANT DISCLAIMERS")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_CRITICAL
    run.font.name = FONT_FAMILY

    disclaimers = [
        ("Source Attribution & Verification:", "This report uses a 4-tier source labeling system. Claims labeled 'Verified Source' have been cross-checked against source material. Claims labeled 'Retrieved Source' come from the curated knowledge base. Claims labeled 'Live Research' come from controlled web searches of curated regulatory sites. Claims labeled 'Model Inference' come from the AI model's training data and have NOT been verified against source material — they require independent review by qualified compliance counsel before being acted upon."),
        ("AI Hallucination Warning:", "This report was generated by an AI model and may contain inaccuracies, including fabricated regulatory citations or misinterpreted requirements. Every finding, citation, and suggested change MUST be independently verified against the source regulation before being acted upon. Pay special attention to claims labeled 'Model Inference' or 'Unverified'."),
        ("Not Legal Advice:", "This document does not constitute legal advice and does not create an attorney-client relationship. It is intended as a starting point for internal review only. Consult qualified healthcare compliance counsel before relying on any content herein."),
        ("PHI Handling:", "No policy text or protected health information is stored by this system. All processing is ephemeral and in-memory only."),
    ]

    for title, body in disclaimers:
        p = cell.add_paragraph()
        run_t = p.add_run(title + " ")
        run_t.bold = True
        run_t.font.size = Pt(9)
        run_t.font.color.rgb = COLOR_CRITICAL
        run_t.font.name = FONT_FAMILY
        run_b = p.add_run(body)
        run_b.font.size = Pt(9)
        run_b.font.color.rgb = COLOR_BLACK
        run_b.font.name = FONT_FAMILY


def _setup_document(doc: Document):
    """Set default document styling."""
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_FAMILY
    font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


# ──────────────────────────────────────────────
# SOURCE ATTRIBUTION HELPERS
# ──────────────────────────────────────────────

def _add_source_attribution_badge(doc, attribution: SourceAttribution):
    """
    Add a source attribution badge/label to the document.
    Shows source type, verification status, source name, and any warnings.
    """
    source_label = SOURCE_TYPE_LABELS.get(attribution.source_type, "Unknown")
    source_color = SOURCE_TYPE_COLORS.get(attribution.source_type, COLOR_GRAY)
    bg_hex = SOURCE_TYPE_BG_HEX.get(attribution.source_type, "F5F5F5")

    verif_label = VERIFICATION_STATUS_LABELS.get(attribution.verification_status, "Unknown")
    verif_color = VERIFICATION_STATUS_COLORS.get(attribution.verification_status, COLOR_GRAY)

    # Build the badge as a small shaded table
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_shading(cell, bg_hex)

    # Line 1: Source type + Verification status
    p = cell.paragraphs[0]
    run_src = p.add_run(f"[{source_label}]")
    run_src.bold = True
    run_src.font.size = Pt(8)
    run_src.font.color.rgb = source_color
    run_src.font.name = FONT_FAMILY

    run_sep = p.add_run("  ")
    run_sep.font.size = Pt(8)

    run_ver = p.add_run(f"[{verif_label}]")
    run_ver.bold = True
    run_ver.font.size = Pt(8)
    run_ver.font.color.rgb = verif_color
    run_ver.font.name = FONT_FAMILY

    # Line 2: Source name and citation if available
    if attribution.source_name or attribution.source_citation:
        p2 = cell.add_paragraph()
        if attribution.source_name:
            run_name = p2.add_run(f"Source: {attribution.source_name}")
            run_name.font.size = Pt(8)
            run_name.font.color.rgb = COLOR_GRAY
            run_name.font.name = FONT_FAMILY
        if attribution.source_citation:
            run_cite = p2.add_run(f"  |  Citation: {attribution.source_citation}")
            run_cite.font.size = Pt(8)
            run_cite.font.color.rgb = COLOR_LIGHT_GRAY
            run_cite.font.name = FONT_FAMILY
        if attribution.source_url:
            run_url = p2.add_run(f"  |  URL: {attribution.source_url}")
            run_url.font.size = Pt(7)
            run_url.font.color.rgb = COLOR_RETRIEVED_SOURCE
            run_url.font.name = FONT_FAMILY

    # Line 3: Supporting evidence (retrieved text) if available
    if attribution.retrieved_text:
        p_evidence = cell.add_paragraph()
        run_ev_label = p_evidence.add_run("Supporting Evidence: ")
        run_ev_label.bold = True
        run_ev_label.font.size = Pt(7)
        run_ev_label.font.color.rgb = COLOR_RETRIEVED_SOURCE
        run_ev_label.font.name = FONT_FAMILY
        # Truncate long evidence to 500 chars in the badge
        evidence_text = attribution.retrieved_text[:500]
        if len(attribution.retrieved_text) > 500:
            evidence_text += "..."
        run_ev_text = p_evidence.add_run(evidence_text)
        run_ev_text.italic = True
        run_ev_text.font.size = Pt(7)
        run_ev_text.font.color.rgb = COLOR_GRAY
        run_ev_text.font.name = FONT_FAMILY

    # Line 4: Confidence score
    if attribution.confidence > 0:
        p4 = cell.add_paragraph()
        run_conf = p4.add_run(f"Confidence: {attribution.confidence:.0%}")
        run_conf.font.size = Pt(7)
        run_conf.font.color.rgb = COLOR_LIGHT_GRAY
        run_conf.font.name = FONT_FAMILY

    # Line 5: Warning if present
    if attribution.warning:
        p_warn = cell.add_paragraph()
        run_warn = p_warn.add_run(f"WARNING: {attribution.warning}")
        run_warn.bold = True
        run_warn.font.size = Pt(8)
        run_warn.font.color.rgb = COLOR_MODEL_INFERENCE
        run_warn.font.name = FONT_FAMILY


def _add_source_attribution_inline(p, attribution: SourceAttribution):
    """
    Add a compact inline source attribution label to an existing paragraph.
    Used for redline changes and checklist items where space is limited.
    """
    source_label = SOURCE_TYPE_LABELS.get(attribution.source_type, "Unknown")
    source_color = SOURCE_TYPE_COLORS.get(attribution.source_type, COLOR_GRAY)
    verif_label = VERIFICATION_STATUS_LABELS.get(attribution.verification_status, "Unknown")

    run_sep = p.add_run("  ")
    run_sep.font.size = Pt(8)

    run_src = p.add_run(f"[{source_label}]")
    run_src.bold = True
    run_src.font.size = Pt(8)
    run_src.font.color.rgb = source_color
    run_src.font.name = FONT_FAMILY

    run_sep2 = p.add_run(" ")
    run_sep2.font.size = Pt(8)

    run_ver = p.add_run(f"[{verif_label}]")
    run_ver.font.size = Pt(8)
    verif_color = VERIFICATION_STATUS_COLORS.get(attribution.verification_status, COLOR_GRAY)
    run_ver.font.color.rgb = verif_color
    run_ver.font.name = FONT_FAMILY

    if attribution.source_name:
        run_name = p.add_run(f" ({attribution.source_name})")
        run_name.font.size = Pt(7)
        run_name.font.color.rgb = COLOR_LIGHT_GRAY
        run_name.font.name = FONT_FAMILY


def _add_source_legend(doc):
    """
    Add the 4-tier Source Attribution Legend explaining the labeling system.
    This appears at the beginning of the document so readers understand
    what each label means.
    """
    _add_styled_paragraph(doc, "SOURCE ATTRIBUTION LEGEND", bold=True, size=12,
                          color=COLOR_DARK_NAVY, space_before=160, space_after=80)

    _add_styled_paragraph(doc, "Every claim in this report carries a source attribution label indicating where the information came from and whether it has been verified. This is the core of the no-hallucination architecture: if a claim cannot be verified from source material, it is clearly labeled as model inference and flagged for independent review.",
                          size=9, space_after=80)

    legend_items = [
        (SourceType.verified_source, "Verified Source",
         "The claim has been cross-checked against source material in the curated knowledge base and confirmed. This is the highest confidence level."),
        (SourceType.retrieved_source, "Retrieved Source",
         "The claim is supported by material retrieved from the curated compliance knowledge base (HIPAA, OCR guidance, OIG guidance, etc.) but has not been individually cross-verified. High confidence but should still be confirmed."),
        (SourceType.live_research, "Live Research",
         "The claim is based on controlled live research from curated regulatory sources (HHS.gov, Federal Register, OCR, CMS, OIG). Current but should be verified against primary source documents before reliance."),
        (SourceType.model_knowledge, "Model Inference",
         "The claim comes from the AI model's training data and has NOT been verified against any loaded source material. This is the lowest confidence level. Such claims MUST be independently verified by qualified compliance counsel before being relied upon."),
    ]

    for source_type, label, description in legend_items:
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        bg_hex = SOURCE_TYPE_BG_HEX.get(source_type, "F5F5F5")
        _set_cell_shading(cell, bg_hex)

        p = cell.paragraphs[0]
        run_label = p.add_run(f"[{label}]")
        run_label.bold = True
        run_label.font.size = Pt(9)
        run_label.font.color.rgb = SOURCE_TYPE_COLORS.get(source_type, COLOR_GRAY)
        run_label.font.name = FONT_FAMILY

        p2 = cell.add_paragraph()
        run_desc = p2.add_run(description)
        run_desc.font.size = Pt(8)
        run_desc.font.color.rgb = COLOR_BLACK
        run_desc.font.name = FONT_FAMILY

    # Verification status legend
    _add_styled_paragraph(doc, "", size=4, space_after=40)
    _add_styled_paragraph(doc, "Verification Status Indicators:", bold=True, size=10,
                          color=COLOR_DARK_NAVY, space_after=40)

    verif_items = [
        (VerificationStatus.verified, "Verified — Confirmed against source material with supporting evidence"),
        (VerificationStatus.partially_verified, "Partially Verified — Some supporting evidence found, but not an exact match"),
        (VerificationStatus.unverified, "Unverified — No supporting evidence found in the source material. Requires independent review."),
        (VerificationStatus.contradicted, "Contradicted — Source material contradicts this claim. Do not rely on it."),
    ]

    for status, description in verif_items:
        p = doc.add_paragraph()
        run_label = p.add_run(f"  {VERIFICATION_STATUS_LABELS.get(status, '')}: ")
        run_label.bold = True
        run_label.font.size = Pt(8)
        run_label.font.color.rgb = VERIFICATION_STATUS_COLORS.get(status, COLOR_GRAY)
        run_label.font.name = FONT_FAMILY
        run_desc = p.add_run(description)
        run_desc.font.size = Pt(8)
        run_desc.font.color.rgb = COLOR_BLACK
        run_desc.font.name = FONT_FAMILY


def _add_sources_used_section(doc, source_names: Optional[List[str]], live_research_used: bool,
                                verification_summary: Optional[str] = None):
    """
    Add a section showing which knowledge base sources were used and whether
    live research contributed to the output.
    """
    if not source_names and not live_research_used and not verification_summary:
        return

    _add_styled_paragraph(doc, "Sources Used", bold=True, size=10,
                          color=COLOR_DARK_NAVY, space_before=80, space_after=40)

    if source_names:
        p = doc.add_paragraph()
        run_label = p.add_run("Knowledge Base Sources: ")
        run_label.bold = True
        run_label.font.size = Pt(9)
        run_label.font.color.rgb = COLOR_RETRIEVED_SOURCE
        run_label.font.name = FONT_FAMILY
        run_val = p.add_run(", ".join(source_names))
        run_val.font.size = Pt(9)
        run_val.font.color.rgb = COLOR_GRAY
        run_val.font.name = FONT_FAMILY

    if live_research_used:
        p = doc.add_paragraph()
        run_label = p.add_run("Live Research: ")
        run_label.bold = True
        run_label.font.size = Pt(9)
        run_label.font.color.rgb = COLOR_LIVE_RESEARCH
        run_label.font.name = FONT_FAMILY
        run_val = p.add_run("Yes — controlled research from curated regulatory sources (HHS.gov, Federal Register, OCR, CMS, OIG)")
        run_val.font.size = Pt(9)
        run_val.font.color.rgb = COLOR_GRAY
        run_val.font.name = FONT_FAMILY

    if verification_summary:
        _add_styled_paragraph(doc, verification_summary, italic=True, size=9,
                              color=COLOR_GRAY, space_before=20, space_after=40)


# ──────────────────────────────────────────────
# COMPLETE ACTION PACKAGE EXPORT
# ──────────────────────────────────────────────

def generate_action_package_docx(package: ComplianceActionPackage, file_name: Optional[str] = None) -> bytes:
    """Generate a comprehensive .docx containing the full Compliance Action Package with source attribution."""
    doc = Document()
    _setup_document(doc)

    # ── Cover Page ──
    _build_cover_page(doc, package, file_name)

    # ── Table of Contents placeholder ──
    _add_styled_paragraph(doc, "TABLE OF CONTENTS", bold=True, size=16, color=COLOR_DARK_NAVY,
                          space_before=100, space_after=80)
    toc_items = []
    section_num = 1
    toc_items.append(f"{section_num}. Source Attribution Legend")
    section_num += 1
    if package.gap_analysis:
        toc_items.append(f"{section_num}. Gap Analysis Findings")
        section_num += 1
    if package.rewritten_policy:
        toc_items.append(f"{section_num}. Rewritten Policy")
        section_num += 1
    if package.redline_changes:
        toc_items.append(f"{section_num}. Redline Document (Tracked Changes)")
        section_num += 1
    if package.remediation_plan:
        toc_items.append(f"{section_num}. 90-Day Remediation Plan")
        section_num += 1
    if package.board_summary:
        toc_items.append(f"{section_num}. Board-Ready Executive Summary")
        section_num += 1
    if package.implementation_checklist:
        toc_items.append(f"{section_num}. Implementation Checklist")
        section_num += 1
    toc_items.append(f"{section_num}. Source Verification Summary")
    section_num += 1
    for item in toc_items:
        _add_styled_paragraph(doc, item, size=11, color=COLOR_GRAY, space_after=40)

    doc.add_page_break()

    # ── Section: Source Attribution Legend ──
    _add_source_legend(doc)

    # ── Section: Gap Analysis ──
    if package.gap_analysis:
        _build_gap_analysis_section(doc, package.gap_analysis, standalone=False)

    # ── Section: Rewritten Policy ──
    if package.rewritten_policy:
        _build_rewritten_policy_section(doc, package.rewritten_policy)

    # ── Section: Redline ──
    if package.redline_changes:
        _build_redline_section(doc, package.redline_changes)

    # ── Section: Remediation Plan ──
    if package.remediation_plan:
        _build_remediation_plan_section(doc, package.remediation_plan)

    # ── Section: Board Summary ──
    if package.board_summary:
        _build_board_summary_section(doc, package.board_summary)

    # ── Section: Implementation Checklist ──
    if package.implementation_checklist:
        _build_checklist_section(doc, package.implementation_checklist)

    # ── Section: Overall Source Verification Summary ──
    _build_verification_summary_section(doc, package)

    # ── Acknowledgment ──
    _build_acknowledgment_section(doc)

    # ── Footer ──
    _build_footer_section(doc)

    # Export to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ── Cover Page ──

def _build_cover_page(doc: Document, package: ComplianceActionPackage, file_name: Optional[str]):
    _add_styled_paragraph(doc, "", size=24, space_after=200)
    _add_styled_paragraph(doc, "COMPLIANCE ACTION PACKAGE", bold=True, size=28,
                          color=COLOR_DARK_NAVY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=80)
    _add_styled_paragraph(doc, "Source-Grounded Regulatory Compliance Deliverable", bold=False, size=14,
                          color=COLOR_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=100)

    _add_horizontal_rule(doc)

    _add_styled_paragraph(doc, "", size=6, space_after=60)
    info_data = [
        ("Policy Document:", file_name or "Uploaded Policy"),
        ("Policy Type:", package.policy_type),
        ("Date of Analysis:", datetime.now().strftime("%B %d, %Y")),
        ("Package ID:", package.package_id),
        ("Outputs Generated:", ", ".join(package.completed_outputs)),
    ]
    if package.gap_analysis:
        info_data.append(("Total Findings:", f"{package.gap_analysis.critical_count + package.gap_analysis.gap_count + package.gap_analysis.partial_count} issues identified"))
        info_data.append(("Regulations Applied:", f"{len(package.gap_analysis.regulations_applied)} regulatory frameworks"))

    # Source attribution info on cover page
    if package.kb_sources_used:
        info_data.append(("Knowledge Base Sources:", f"{len(package.kb_sources_used)} source(s) used"))
    if package.live_research_used:
        info_data.append(("Live Research:", "Used — curated regulatory sources searched"))
    if package.unverified_claim_count is not None and package.unverified_claim_count > 0:
        info_data.append(("Unverified Claims:", f"{package.unverified_claim_count} — require independent review"))

    for label, value in info_data:
        p = doc.add_paragraph()
        run_label = p.add_run(label + "  ")
        run_label.bold = True
        run_label.font.size = Pt(10)
        run_label.font.color.rgb = COLOR_GRAY
        run_label.font.name = FONT_FAMILY
        run_value = p.add_run(value)
        run_value.font.size = Pt(10)
        run_value.font.color.rgb = COLOR_BLACK
        run_value.font.name = FONT_FAMILY

    # KB Sources list on cover page
    if package.kb_sources_used:
        _add_styled_paragraph(doc, "", size=4, space_after=20)
        _add_styled_paragraph(doc, "Knowledge Base Sources Used in This Analysis:", bold=True, size=9,
                              color=COLOR_RETRIEVED_SOURCE, space_after=40)
        for src in package.kb_sources_used:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(src)
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_GRAY
            run.font.name = FONT_FAMILY

    _add_styled_paragraph(doc, "", size=6, space_after=200)
    _add_disclaimer_box(doc)
    doc.add_page_break()


# ── Gap Analysis Section ──

def _build_gap_analysis_section(doc: Document, result: AnalysisResult,
                                 file_name: Optional[str] = None,
                                 standalone: bool = True):
    # Map our 4-tier risk model onto the template's 3-tier vocabulary.
    # Multiple source levels can map to the same target bucket so no row is silently dropped.
    severity_groups = [
        ("Critical", [RiskLevel.critical]),
        ("Moderate", [RiskLevel.high]),
        ("Minor", [RiskLevel.moderate, RiskLevel.low]),
    ]
    grouped = {label: [r for r in result.gap_table if r.risk_level in levels]
               for label, levels in severity_groups}
    counts = {label: len(rows) for label, rows in grouped.items()}
    total_issues = sum(counts.values())

    if standalone:
        # Title
        _add_styled_paragraph(doc, "Policy Gap Analysis Report", bold=True, size=22,
                              color=COLOR_DARK_NAVY, space_after=120)

        # Header metadata
        _add_styled_paragraph(doc, f"Policy: {file_name or 'Uploaded Policy'}",
                              size=11, space_after=40)
        _add_styled_paragraph(doc, f"Generated: {datetime.now().strftime('%B %d, %Y')}",
                              size=11, space_after=40)
        _add_styled_paragraph(doc, f"Policy Type: {result.policy_type}",
                              size=11, space_after=160)
    else:
        # Inside the action package: just a section heading, no duplicate metadata
        _add_styled_paragraph(doc, "Gap Analysis Findings", bold=True, size=18,
                              color=COLOR_DARK_NAVY, space_before=100, space_after=120)

    # Summary of Findings
    _add_styled_paragraph(doc, "Summary of Findings", bold=True, size=14,
                          color=COLOR_DARK_NAVY, space_after=80)
    _add_styled_paragraph(doc, f"Total Issues Identified: {total_issues}", size=11, space_after=40)
    _add_styled_paragraph(doc, f"Critical: {counts['Critical']}", size=11, space_after=20)
    _add_styled_paragraph(doc, f"Moderate: {counts['Moderate']}", size=11, space_after=20)
    _add_styled_paragraph(doc, f"Minor: {counts['Minor']}", size=11, space_after=160)

    # Executive Summary
    _add_styled_paragraph(doc, "Executive Summary", bold=True, size=14,
                          color=COLOR_DARK_NAVY, space_after=80)
    _add_styled_paragraph(doc, result.audit_ready_summary, size=11, space_after=160)

    # Findings sections
    for label, _level in severity_groups:
        rows = grouped[label]
        if not rows:
            continue
        _add_styled_paragraph(doc, f"{label} Findings ({len(rows)})", bold=True, size=14,
                              color=COLOR_DARK_NAVY, space_before=80, space_after=80)
        for idx, row in enumerate(rows, 1):
            # Numbered title: "1. Clause"
            p = doc.add_paragraph()
            run_num = p.add_run(f"{idx}. ")
            run_num.bold = True
            run_num.font.size = Pt(11)
            run_num.font.color.rgb = COLOR_BLACK
            run_num.font.name = FONT_FAMILY
            run_clause = p.add_run(row.clause)
            run_clause.bold = True
            run_clause.font.size = Pt(11)
            run_clause.font.color.rgb = COLOR_BLACK
            run_clause.font.name = FONT_FAMILY
            p.paragraph_format.space_after = Pt(3)

            # Applicable Regulations
            regs = ", ".join(row.regulations) if row.regulations else row.citation
            p = doc.add_paragraph()
            run_l = p.add_run("Applicable Regulations: ")
            run_l.bold = True
            run_l.font.size = Pt(10)
            run_l.font.color.rgb = COLOR_BLACK
            run_l.font.name = FONT_FAMILY
            run_v = p.add_run(regs)
            run_v.font.size = Pt(10)
            run_v.font.color.rgb = COLOR_BLACK
            run_v.font.name = FONT_FAMILY
            p.paragraph_format.space_after = Pt(3)

            # Finding
            p = doc.add_paragraph()
            run_l = p.add_run("Finding: ")
            run_l.bold = True
            run_l.font.size = Pt(10)
            run_l.font.color.rgb = COLOR_BLACK
            run_l.font.name = FONT_FAMILY
            run_v = p.add_run(row.finding)
            run_v.font.size = Pt(10)
            run_v.font.color.rgb = COLOR_BLACK
            run_v.font.name = FONT_FAMILY
            p.paragraph_format.space_after = Pt(3)

            # Suggested Language
            p = doc.add_paragraph()
            run_l = p.add_run("Suggested Language: ")
            run_l.bold = True
            run_l.font.size = Pt(10)
            run_l.font.color.rgb = COLOR_BLACK
            run_l.font.name = FONT_FAMILY
            run_v = p.add_run(row.suggested_language)
            run_v.font.size = Pt(10)
            run_v.font.color.rgb = COLOR_BLACK
            run_v.font.name = FONT_FAMILY
            p.paragraph_format.space_after = Pt(3)

            # Citation
            p = doc.add_paragraph()
            run_l = p.add_run("Citation: ")
            run_l.bold = True
            run_l.font.size = Pt(10)
            run_l.font.color.rgb = COLOR_BLACK
            run_l.font.name = FONT_FAMILY
            run_v = p.add_run(row.citation)
            run_v.font.size = Pt(10)
            run_v.font.color.rgb = COLOR_BLACK
            run_v.font.name = FONT_FAMILY
            p.paragraph_format.space_after = Pt(12)

    # All Regulations Reviewed
    if result.regulations_applied:
        _add_styled_paragraph(doc, "All Regulations Reviewed", bold=True, size=14,
                              color=COLOR_DARK_NAVY, space_before=160, space_after=80)
        for reg in result.regulations_applied:
            _add_styled_paragraph(doc, reg, size=11, space_after=20)


# ── Rewritten Policy Section ──

def _build_rewritten_policy_section(doc: Document, rewritten: RewrittenPolicy):
    _add_styled_paragraph(doc, "3. REWRITTEN POLICY", bold=True, size=16,
                          color=COLOR_DARK_NAVY, space_before=200, space_after=80)
    _add_styled_paragraph(doc, rewritten.policy_title, bold=True, size=14,
                          color=COLOR_ACCENT_ORANGE, space_after=40)
    _add_styled_paragraph(doc, f"Effective Date: {rewritten.effective_date or 'Upon adoption'}", size=10,
                          color=COLOR_GRAY, space_after=40)
    _add_styled_paragraph(doc, f"Version Note: {rewritten.version_note}", italic=True, size=10,
                          color=COLOR_GRAY, space_after=40)
    _add_styled_paragraph(doc, f"Change Summary: {rewritten.change_summary}", size=10, space_after=100)

    # Section-by-section changes
    for section in rewritten.sections:
        _add_styled_paragraph(doc, section.section_title, bold=True, size=12,
                              color=COLOR_DARK_NAVY, space_before=120, space_after=40)

        # Source attribution for this section if available
        if section.source_attribution:
            _add_source_attribution_badge(doc, section.source_attribution)
            _add_styled_paragraph(doc, "", size=2, space_after=20)

        if section.original_text and section.original_text != "NEW SECTION":
            _add_styled_paragraph(doc, "Original Text:", bold=True, size=9,
                                  color=COLOR_HIGH, space_after=20)
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.cell(0, 0)
            _set_cell_shading(cell, "FFF0F0")
            p = cell.paragraphs[0]
            run = p.add_run(section.original_text[:2000])
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_BLACK
            run.font.name = FONT_FAMILY

        _add_styled_paragraph(doc, "Rewritten Text:", bold=True, size=9,
                              color=COLOR_COMPLIANT, space_before=60, space_after=20)
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        _set_cell_shading(cell, "F0FFF4")
        p = cell.paragraphs[0]
        run = p.add_run(section.rewritten_text[:3000])
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_BLACK
        run.font.name = FONT_FAMILY

        if section.changes_summary:
            _add_styled_paragraph(doc, f"Changes: {section.changes_summary}", italic=True, size=9,
                                  color=COLOR_GRAY, space_before=20, space_after=40)
        if section.regulation_refs:
            _add_styled_paragraph(doc, f"Regulations: {', '.join(section.regulation_refs)}", size=9,
                                  color=COLOR_LIGHT_GRAY, space_after=60)

    # Full text
    _add_styled_paragraph(doc, "Complete Rewritten Policy (Ready for Adoption)", bold=True, size=12,
                          color=COLOR_DARK_NAVY, space_before=200, space_after=80)
    _add_styled_paragraph(doc, rewritten.full_text, size=10, space_after=100)

    # Sources used for this section
    _add_sources_used_section(
        doc,
        source_names=rewritten.retrieved_sources_used,
        live_research_used=rewritten.live_research_used,
    )


# ── Redline Section ──

def _build_redline_section(doc: Document, changes: List[RedlineChange]):
    _add_styled_paragraph(doc, "4. REDLINE DOCUMENT (TRACKED CHANGES)", bold=True, size=16,
                          color=COLOR_DARK_NAVY, space_before=200, space_after=80)
    _add_styled_paragraph(doc, f"Total Changes Tracked: {len(changes)}", bold=True, size=11, space_after=100)

    added = sum(1 for c in changes if c.type == "added")
    removed = sum(1 for c in changes if c.type == "removed")
    modified = sum(1 for c in changes if c.type == "modified")

    _add_styled_paragraph(doc, f"Added: {added}  |  Removed: {removed}  |  Modified: {modified}",
                          size=10, color=COLOR_GRAY, space_after=80)

    current_section = ""
    for change in changes:
        # Section header
        if change.section and change.section != current_section:
            current_section = change.section
            _add_styled_paragraph(doc, current_section, bold=True, size=11,
                                  color=COLOR_DARK_NAVY, space_before=100, space_after=40)

        if change.type == "added":
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            _set_cell_shading(cell, "E8F5E9")
            p = cell.paragraphs[0]
            run_label = p.add_run("[ADDED] ")
            run_label.bold = True
            run_label.font.size = Pt(9)
            run_label.font.color.rgb = COLOR_COMPLIANT
            run_label.font.name = FONT_FAMILY
            run_text = p.add_run(change.revised_text or "")
            run_text.font.size = Pt(9)
            run_text.font.color.rgb = COLOR_BLACK
            run_text.font.name = FONT_FAMILY

            # Source attribution inline for this change
            if change.source_attribution:
                _add_source_attribution_inline(p, change.source_attribution)

        elif change.type == "removed":
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            _set_cell_shading(cell, "FFEBEE")
            p = cell.paragraphs[0]
            run_label = p.add_run("[REMOVED] ")
            run_label.bold = True
            run_label.font.size = Pt(9)
            run_label.font.color.rgb = COLOR_CRITICAL
            run_label.font.name = FONT_FAMILY
            run_text = p.add_run(change.original_text or "")
            run_text.font.size = Pt(9)
            run_text.font.color.rgb = COLOR_BLACK
            run_text.font.name = FONT_FAMILY
            run_text.font.strike = True

        elif change.type == "modified":
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            _set_cell_shading(cell, "FFF8E1")
            p = cell.paragraphs[0]
            run_label = p.add_run("[MODIFIED] ")
            run_label.bold = True
            run_label.font.size = Pt(9)
            run_label.font.color.rgb = COLOR_MODERATE
            run_label.font.name = FONT_FAMILY
            if change.original_text:
                run_old = p.add_run(f"\nOriginal: {change.original_text[:500]}")
                run_old.font.size = Pt(9)
                run_old.font.color.rgb = COLOR_HIGH
                run_old.font.name = FONT_FAMILY
                run_old.font.strike = True
            if change.revised_text:
                run_new = p.add_run(f"\nRevised: {change.revised_text[:500]}")
                run_new.font.size = Pt(9)
                run_new.font.color.rgb = COLOR_COMPLIANT
                run_new.font.name = FONT_FAMILY

            # Source attribution inline for modified changes
            if change.source_attribution:
                _add_source_attribution_inline(p, change.source_attribution)

        if change.regulation_ref:
            _add_styled_paragraph(doc, f"Regulation: {change.regulation_ref}", size=8,
                                  color=COLOR_LIGHT_GRAY, space_before=10, space_after=40)

        # Full source attribution badge for redline changes with attribution
        if change.source_attribution:
            _add_source_attribution_badge(doc, change.source_attribution)
            _add_styled_paragraph(doc, "", size=2, space_after=20)


# ── Remediation Plan Section ──

def _build_remediation_plan_section(doc: Document, plan: RemediationPlan):
    _add_styled_paragraph(doc, "6. 90-DAY REMEDIATION PLAN", bold=True, size=16,
                          color=COLOR_DARK_NAVY, space_before=200, space_after=80)
    _add_styled_paragraph(doc, plan.plan_title, bold=True, size=13,
                          color=COLOR_ACCENT_ORANGE, space_after=60)
    _add_styled_paragraph(doc, f"Total Tasks: {plan.total_tasks}  |  Critical (First 30 Days): {plan.critical_tasks_first_30}",
                          size=11, space_after=80)

    for phase in plan.phases:
        phase_color = {1: COLOR_CRITICAL, 2: COLOR_HIGH, 3: COLOR_MODERATE}
        color = phase_color.get(phase.phase_number, COLOR_MODERATE)

        _add_styled_paragraph(doc, f"Phase {phase.phase_number}: {phase.phase_name} ({phase.time_range})",
                              bold=True, size=12, color=color, space_before=120, space_after=40)
        _add_styled_paragraph(doc, phase.objective, size=10, space_after=60)

        for task in phase.tasks:
            task_color = {"critical": COLOR_CRITICAL, "high": COLOR_HIGH, "moderate": COLOR_MODERATE}
            t_color = task_color.get(task.risk_level, COLOR_MODERATE)

            p = doc.add_paragraph()
            run_id = p.add_run(f"{task.task_id} ")
            run_id.bold = True
            run_id.font.size = Pt(10)
            run_id.font.color.rgb = t_color
            run_id.font.name = FONT_FAMILY
            run_title = p.add_run(task.title)
            run_title.bold = True
            run_title.font.size = Pt(10)
            run_title.font.color.rgb = COLOR_BLACK
            run_title.font.name = FONT_FAMILY

            # Source attribution inline for tasks
            if task.source_attribution:
                _add_source_attribution_inline(p, task.source_attribution)

            _add_styled_paragraph(doc, f"Responsible: {task.responsible_party}  |  Deliverable: {task.deliverable}",
                                  size=9, color=COLOR_GRAY, space_after=20)
            _add_styled_paragraph(doc, task.description, size=10, space_after=20)
            _add_styled_paragraph(doc, f"Verification: {task.verification_method}", italic=True, size=9,
                                  color=COLOR_GRAY, space_after=20)

            # Full source attribution badge for tasks
            if task.source_attribution:
                _add_source_attribution_badge(doc, task.source_attribution)
                _add_styled_paragraph(doc, "", size=2, space_after=20)

    _add_styled_paragraph(doc, "Success Criteria", bold=True, size=12,
                          color=COLOR_DARK_NAVY, space_before=100, space_after=40)
    _add_styled_paragraph(doc, plan.success_criteria, size=10, space_after=60)

    _add_styled_paragraph(doc, "Resource Requirements", bold=True, size=12,
                          color=COLOR_DARK_NAVY, space_before=60, space_after=40)
    _add_styled_paragraph(doc, plan.resource_requirements, size=10, space_after=100)

    # Sources used
    _add_sources_used_section(
        doc,
        source_names=plan.retrieved_sources_used,
        live_research_used=plan.live_research_used,
    )


# ── Board Summary Section ──

def _build_board_summary_section(doc: Document, summary: BoardSummary):
    _add_styled_paragraph(doc, "7. BOARD-READY EXECUTIVE SUMMARY", bold=True, size=16,
                          color=COLOR_DARK_NAVY, space_before=200, space_after=80)

    # Headline in a prominent box
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_shading(cell, "1A1A2E")
    p = cell.paragraphs[0]
    run = p.add_run(summary.headline)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_WHITE
    run.font.name = FONT_FAMILY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_styled_paragraph(doc, "", size=6, space_after=40)

    _add_styled_paragraph(doc, f"Overall Status: {summary.overall_status}", bold=True, size=12, space_after=60)
    _add_styled_paragraph(doc, summary.risk_summary, size=11, space_after=80)

    _add_styled_paragraph(doc, "Key Findings for Board Awareness", bold=True, size=11,
                          color=COLOR_DARK_NAVY, space_after=40)
    for finding in summary.key_findings:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(finding)
        run.font.size = Pt(10)
        run.font.name = FONT_FAMILY

    _add_styled_paragraph(doc, "Regulatory Exposure", bold=True, size=11,
                          color=COLOR_HIGH, space_before=60, space_after=40)
    _add_styled_paragraph(doc, summary.regulatory_exposure, size=10, space_after=60)

    _add_styled_paragraph(doc, "Remediation Status", bold=True, size=11,
                          color=COLOR_DARK_NAVY, space_before=40, space_after=40)
    _add_styled_paragraph(doc, summary.remediation_status, size=10, space_after=60)

    _add_styled_paragraph(doc, "Recommended Board Actions", bold=True, size=11,
                          color=COLOR_DARK_NAVY, space_before=40, space_after=40)
    for action in summary.recommended_actions:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(action)
        run.font.size = Pt(10)
        run.font.name = FONT_FAMILY

    if summary.budget_impact:
        _add_styled_paragraph(doc, f"Budget Impact: {summary.budget_impact}", bold=True, size=10,
                              color=COLOR_ACCENT_ORANGE, space_before=60, space_after=40)

    if summary.next_review_date:
        _add_styled_paragraph(doc, f"Next Review: {summary.next_review_date}", size=10,
                              color=COLOR_GRAY, space_after=40)

    _add_styled_paragraph(doc, f"Prepared by: {summary.prepared_by or 'Chief Compliance Officer'}  |  Date: {summary.prepared_date or datetime.now().strftime('%B %d, %Y')}",
                          size=9, color=COLOR_LIGHT_GRAY, space_after=60)

    # Sources used for board summary
    _add_sources_used_section(
        doc,
        source_names=summary.retrieved_sources_used,
        live_research_used=summary.live_research_used,
    )


# ── Implementation Checklist Section ──

def _build_checklist_section(doc: Document, checklist: ImplementationChecklist):
    _add_styled_paragraph(doc, "8. IMPLEMENTATION CHECKLIST", bold=True, size=16,
                          color=COLOR_DARK_NAVY, space_before=200, space_after=80)
    _add_styled_paragraph(doc, f"Total Items: {checklist.total_items}  |  Critical: {checklist.critical_items}",
                          bold=True, size=11, space_after=40)
    _add_styled_paragraph(doc, checklist.completion_timeline, size=10, color=COLOR_GRAY, space_after=80)

    # Group by category
    by_category = {}
    for item in checklist.items:
        by_category.setdefault(item.category, []).append(item)

    for category, items in by_category.items():
        _add_styled_paragraph(doc, category, bold=True, size=12,
                              color=COLOR_DARK_NAVY, space_before=100, space_after=40)

        for item in items:
            priority_color = {"critical": COLOR_CRITICAL, "high": COLOR_HIGH, "moderate": COLOR_MODERATE, "low": COLOR_LOW}
            color = priority_color.get(item.priority, COLOR_MODERATE)

            p = doc.add_paragraph()
            run_check = p.add_run("[ ] ")
            run_check.font.size = Pt(10)
            run_check.font.name = FONT_FAMILY
            run_id = p.add_run(f"{item.item_id} ")
            run_id.bold = True
            run_id.font.size = Pt(9)
            run_id.font.color.rgb = color
            run_id.font.name = FONT_FAMILY
            run_action = p.add_run(item.action)
            run_action.font.size = Pt(10)
            run_action.font.color.rgb = COLOR_BLACK
            run_action.font.name = FONT_FAMILY

            # Source attribution inline for checklist items
            if item.source_attribution:
                _add_source_attribution_inline(p, item.source_attribution)

            _add_styled_paragraph(doc, f"Responsible: {item.responsible_role}  |  Deadline: {item.deadline}  |  Priority: {item.priority.upper()}",
                                  size=9, color=COLOR_GRAY, space_after=10)
            _add_styled_paragraph(doc, f"Regulation: {item.regulation_ref}", size=9, color=COLOR_LIGHT_GRAY, space_after=10)
            _add_styled_paragraph(doc, f"Verification: {item.verification}", italic=True, size=9,
                                  color=COLOR_GRAY, space_after=10)
            _add_styled_paragraph(doc, f"Evidence Needed: {item.evidence_needed}", italic=True, size=9,
                                  color=COLOR_GRAY, space_after=20)

            # Full source attribution badge for checklist items
            if item.source_attribution:
                _add_source_attribution_badge(doc, item.source_attribution)
                _add_styled_paragraph(doc, "", size=2, space_after=20)

    # Sources used
    _add_sources_used_section(
        doc,
        source_names=checklist.retrieved_sources_used,
        live_research_used=checklist.live_research_used,
    )


# ── Overall Source Verification Summary Section ──

def _build_verification_summary_section(doc: Document, package: ComplianceActionPackage):
    """
    Build the overall source verification summary section at the end of the document.
    This gives readers a consolidated view of what was verified and what wasn't.
    """
    _add_styled_paragraph(doc, "9. SOURCE VERIFICATION SUMMARY", bold=True, size=16,
                          color=COLOR_DARK_NAVY, space_before=200, space_after=80)

    _add_styled_paragraph(doc, "This section provides a consolidated summary of source attribution and verification across all outputs in this Compliance Action Package. It is critical that all unverified claims are independently reviewed by qualified compliance counsel before being relied upon.",
                          size=10, space_after=80)

    # Overall verification status
    if package.verification_overall:
        # Determine the color based on whether there are unverified claims
        if package.unverified_claim_count and package.unverified_claim_count > 0:
            verif_bg = "FFF8E1"  # Amber background
            verif_color = COLOR_MODEL_INFERENCE
        elif package.kb_sources_used:
            verif_bg = "E8F5E9"  # Green background
            verif_color = COLOR_COMPLIANT
        else:
            verif_bg = "FFF8E1"  # Amber background — model only
            verif_color = COLOR_MODEL_INFERENCE

        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        _set_cell_shading(cell, verif_bg)
        p = cell.paragraphs[0]
        run = p.add_run(package.verification_overall)
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_BLACK
        run.font.name = FONT_FAMILY

    # KB Sources Used
    if package.kb_sources_used:
        _add_styled_paragraph(doc, "", size=4, space_after=40)
        _add_styled_paragraph(doc, "Knowledge Base Sources Used Across All Outputs:", bold=True, size=11,
                              color=COLOR_RETRIEVED_SOURCE, space_after=40)
        for src in package.kb_sources_used:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(src)
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_GRAY
            run.font.name = FONT_FAMILY

    # Live Research
    if package.live_research_used:
        _add_styled_paragraph(doc, "", size=4, space_after=40)
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        _set_cell_shading(cell, "F3E5F5")  # Light purple
        p = cell.paragraphs[0]
        run_label = p.add_run("Live Research Was Used")
        run_label.bold = True
        run_label.font.size = Pt(10)
        run_label.font.color.rgb = COLOR_LIVE_RESEARCH
        run_label.font.name = FONT_FAMILY
        p2 = cell.add_paragraph()
        run_desc = p2.add_run("Some findings in this package were augmented with controlled live research from curated regulatory sources (HHS.gov, Federal Register, OCR, CMS, OIG). These results are current but should be verified against primary source documents.")
        run_desc.font.size = Pt(9)
        run_desc.font.color.rgb = COLOR_BLACK
        run_desc.font.name = FONT_FAMILY

    # Unverified claims count
    if package.unverified_claim_count is not None and package.unverified_claim_count > 0:
        _add_styled_paragraph(doc, "", size=4, space_after=40)
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        _set_cell_shading(cell, "FFF8E1")  # Light amber
        p = cell.paragraphs[0]
        run_label = p.add_run(f"ATTENTION: {package.unverified_claim_count} Unverified Claim(s)")
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_label.font.color.rgb = COLOR_MODEL_INFERENCE
        run_label.font.name = FONT_FAMILY
        p2 = cell.add_paragraph()
        run_desc = p2.add_run(f"{package.unverified_claim_count} claim(s) in this package could not be verified against loaded source material. These claims are based on model inference only and MUST be independently verified by qualified compliance counsel before being relied upon. Do not treat unverified claims as established regulatory requirements without independent confirmation.")
        run_desc.font.size = Pt(9)
        run_desc.font.color.rgb = COLOR_BLACK
        run_desc.font.name = FONT_FAMILY

    # Model-only mode warning
    if not package.kb_sources_used and not package.live_research_used:
        _add_styled_paragraph(doc, "", size=4, space_after=40)
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        _set_cell_shading(cell, "FFEBEE")  # Light red
        p = cell.paragraphs[0]
        run_label = p.add_run("MODEL-ONLY MODE — No Source Material Was Available")
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_label.font.color.rgb = COLOR_CRITICAL
        run_label.font.name = FONT_FAMILY
        p2 = cell.add_paragraph()
        run_desc = p2.add_run("No source material was available in the knowledge base during this analysis. ALL findings in this package are model inference only and MUST be independently verified by qualified compliance counsel. The knowledge base should be populated with regulatory source material before this system can produce source-grounded, verified analysis. Use the /api/kb/seed endpoint to load foundational regulatory content.")
        run_desc.font.size = Pt(9)
        run_desc.font.color.rgb = COLOR_BLACK
        run_desc.font.name = FONT_FAMILY

    # Per-section verification breakdown
    _add_styled_paragraph(doc, "", size=4, space_after=60)
    _add_styled_paragraph(doc, "Per-Output Verification Summary:", bold=True, size=11,
                          color=COLOR_DARK_NAVY, space_after=40)

    # Build a summary table for each output
    output_verification = []
    if package.gap_analysis and package.gap_analysis.source_attributions:
        verified = sum(1 for a in package.gap_analysis.source_attributions if a.verification_status == VerificationStatus.verified)
        unverified = sum(1 for a in package.gap_analysis.source_attributions if a.verification_status in (VerificationStatus.unverified, VerificationStatus.contradicted))
        output_verification.append(("Gap Analysis", len(package.gap_analysis.source_attributions), verified, unverified, package.gap_analysis.live_research_used))

    if package.rewritten_policy and package.rewritten_policy.source_attributions:
        verified = sum(1 for a in package.rewritten_policy.source_attributions if a.verification_status == VerificationStatus.verified)
        unverified = sum(1 for a in package.rewritten_policy.source_attributions if a.verification_status in (VerificationStatus.unverified, VerificationStatus.contradicted))
        output_verification.append(("Rewritten Policy", len(package.rewritten_policy.source_attributions), verified, unverified, package.rewritten_policy.live_research_used))

    if package.remediation_plan and package.remediation_plan.source_attributions:
        verified = sum(1 for a in package.remediation_plan.source_attributions if a.verification_status == VerificationStatus.verified)
        unverified = sum(1 for a in package.remediation_plan.source_attributions if a.verification_status in (VerificationStatus.unverified, VerificationStatus.contradicted))
        output_verification.append(("Remediation Plan", len(package.remediation_plan.source_attributions), verified, unverified, package.remediation_plan.live_research_used))

    if package.board_summary and package.board_summary.source_attributions:
        verified = sum(1 for a in package.board_summary.source_attributions if a.verification_status == VerificationStatus.verified)
        unverified = sum(1 for a in package.board_summary.source_attributions if a.verification_status in (VerificationStatus.unverified, VerificationStatus.contradicted))
        output_verification.append(("Board Summary", len(package.board_summary.source_attributions), verified, unverified, package.board_summary.live_research_used))

    if package.implementation_checklist and package.implementation_checklist.source_attributions:
        verified = sum(1 for a in package.implementation_checklist.source_attributions if a.verification_status == VerificationStatus.verified)
        unverified = sum(1 for a in package.implementation_checklist.source_attributions if a.verification_status in (VerificationStatus.unverified, VerificationStatus.contradicted))
        output_verification.append(("Checklist", len(package.implementation_checklist.source_attributions), verified, unverified, package.implementation_checklist.live_research_used))

    if output_verification:
        # Summary table
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        headers = ["Output", "Total Claims", "Verified", "Unverified", "Live Research"]
        for i, header in enumerate(headers):
            _set_cell_shading(hdr_cells[i], "1A1A2E")
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = COLOR_WHITE
            run.font.name = FONT_FAMILY
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for output_name, total, verified, unverified, live in output_verification:
            row = table.add_row()
            cells = row.cells
            for i, val in enumerate([output_name, str(total), str(verified), str(unverified), "Yes" if live else "No"]):
                p = cells[i].paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(8)
                run.font.name = FONT_FAMILY
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Color unverified count red
                if i == 3 and unverified > 0:
                    run.font.color.rgb = COLOR_MODEL_INFERENCE
                    run.bold = True
                elif i == 4 and live:
                    run.font.color.rgb = COLOR_LIVE_RESEARCH
    else:
        _add_styled_paragraph(doc, "No detailed source attribution data available for individual outputs. This may indicate the knowledge base was not populated during analysis.",
                              size=9, italic=True, color=COLOR_GRAY, space_after=60)


# ── Acknowledgment Section ──

def _build_acknowledgment_section(doc: Document):
    _add_styled_paragraph(doc, "REVIEW AND ACKNOWLEDGMENT", bold=True, size=16,
                          color=COLOR_DARK_NAVY, space_before=200, space_after=80)
    _add_styled_paragraph(doc, "By signing below, the reviewing compliance officer acknowledges receipt and review of this Compliance Action Package, including the source attribution and verification status of each finding. This acknowledgment does not constitute agreement with all findings, but confirms that the report has been received and will be addressed in accordance with organizational compliance procedures. The reviewer specifically acknowledges that claims labeled 'Model Inference' or 'Unverified' require independent verification before being relied upon.",
                          size=10, space_after=120)

    sig_fields = ["Reviewed By (Print Name)", "Title / Department", "Signature", "Date"]
    for field in sig_fields:
        p = doc.add_paragraph()
        run = p.add_run(f"{field}: ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_GRAY
        run.font.name = FONT_FAMILY
        run_line = p.add_run("_" * 50)
        run_line.font.size = Pt(10)
        run_line.font.color.rgb = COLOR_GRAY
        run_line.font.name = FONT_FAMILY


# ── Footer ──

def _build_footer_section(doc: Document):
    _add_styled_paragraph(doc, "", size=6, space_after=40)
    _add_horizontal_rule(doc, "CCCCCC")
    _add_styled_paragraph(doc, "CONFIDENTIAL — For compliance and privacy department use only. Not legal advice. Source-attributed: claims labeled 'Model Inference' require independent verification.",
                          italic=True, size=8, color=COLOR_LIGHT_GRAY, space_after=20)
    _add_styled_paragraph(doc, f"Generated by Policy Gap Analyzer v3.0 (Source-Grounded) | {datetime.now().strftime('%B %d, %Y')}",
                          size=8, color=COLOR_LIGHT_GRAY, space_after=20)
    _add_styled_paragraph(doc, "Built by Andrew Weingarten", size=8, color=COLOR_LIGHT_GRAY)


# ──────────────────────────────────────────────
# Legacy single-report export (backward compatible)
# ──────────────────────────────────────────────

def generate_docx(result: AnalysisResult, file_name: Optional[str] = None,
                   kb_sources_used: Optional[List[str]] = None,
                   live_research_used: bool = False,
                   verification_overall: Optional[str] = None) -> bytes:
    """Generate a clean, template-style .docx gap analysis report and return as bytes.
    Opens with a one-page certificate-style summary (regulations reviewed) before
    the detailed findings -- this used to require a separate certificate download
    for the same information.

    kb_sources_used/live_research_used/verification_overall live on the parent
    ComplianceActionPackage, not on AnalysisResult itself, so they're passed in
    separately -- without them this report had no verification section at all,
    despite the whole point of the app being source-grounded verification."""
    doc = Document()
    _setup_document(doc)
    _build_certificate_content(
        doc,
        policy_type=result.policy_type,
        regulations=result.regulations_applied,
        date_str=datetime.now().strftime("%B %d, %Y"),
    )
    doc.add_page_break()
    _build_gap_analysis_section(doc, result, file_name=file_name)
    doc.add_paragraph()
    _add_sources_used_section(doc, kb_sources_used, live_research_used, verification_overall)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def generate_export(result: AnalysisResult, file_name: Optional[str] = None,
                    export_format: ExportFormat = ExportFormat.docx,
                    kb_sources_used: Optional[List[str]] = None,
                    live_research_used: bool = False,
                    verification_overall: Optional[str] = None) -> tuple[bytes, str]:
    """Generate a single gap analysis export file."""
    safe_name = (file_name or "Policy").rsplit(".", 1)[0].replace(" ", "_").replace("(", "").replace(")", "")[:60]
    date_str = datetime.now().strftime("%Y-%m-%d")
    file_bytes = generate_docx(result, file_name, kb_sources_used, live_research_used, verification_overall)
    filename = f"{safe_name}_Gap_Report_{date_str}.docx"
    return file_bytes, filename


def generate_action_package_export(package: ComplianceActionPackage, file_name: Optional[str] = None,
                                   export_format: ExportFormat = ExportFormat.docx) -> tuple[bytes, str]:
    """Generate the full Compliance Action Package export."""
    safe_name = (file_name or "Policy").rsplit(".", 1)[0].replace(" ", "_").replace("(", "").replace(")", "")[:60]
    date_str = datetime.now().strftime("%Y-%m-%d")
    file_bytes = generate_action_package_docx(package, file_name)
    filename = f"{safe_name}_Compliance_Action_Package_{date_str}.docx"
    return file_bytes, filename


def generate_draft_policy_docx(policy: dict) -> bytes:
    """
    Generate a professional .docx for a drafted policy document.
    `policy` is the DraftedPolicy dict with keys:
      policy_title, effective_date, version, scope, regulations_applied, sections, full_text, drafting_notes
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.2)
        section.right_margin = Cm(2.5)

    # ── Cover / Header ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_p.add_run(policy.get("policy_title", "Policy Document"))
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = COLOR_DARK_NAVY
    run.font.name = FONT_FAMILY
    title_p.paragraph_format.space_after = Pt(4)

    _add_horizontal_rule(doc)

    # Metadata row
    meta_parts = []
    if policy.get("effective_date"):
        meta_parts.append(f"Effective: {policy['effective_date']}")
    if policy.get("version"):
        meta_parts.append(f"Version {policy['version']}")
    meta_parts.append(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    _add_styled_paragraph(doc, "  |  ".join(meta_parts), size=10, color=COLOR_GRAY)

    doc.add_paragraph()

    # Scope
    if policy.get("scope"):
        _add_styled_paragraph(doc, "SCOPE", bold=True, size=10, color=COLOR_ACCENT_ORANGE)
        _add_styled_paragraph(doc, policy["scope"], size=11)
        doc.add_paragraph()

    # Regulations applied
    regs = policy.get("regulations_applied", [])
    if regs:
        _add_styled_paragraph(doc, "REGULATORY FRAMEWORK", bold=True, size=10, color=COLOR_ACCENT_ORANGE)
        for reg in regs:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(reg)
            run.font.size = Pt(10)
            run.font.name = FONT_FAMILY
            run.font.color.rgb = COLOR_DARK_NAVY
        doc.add_paragraph()

    _add_horizontal_rule(doc)
    doc.add_paragraph()

    # Policy sections
    sections = policy.get("sections", [])
    if sections:
        for sec in sections:
            title = sec.get("title", "").strip()
            content = sec.get("content", "").strip()
            if title:
                _add_styled_paragraph(doc, title.upper(), bold=True, size=12,
                                      color=COLOR_DARK_NAVY, space_before=120, space_after=60)
            if content:
                for para_text in content.split("\n\n"):
                    para_text = para_text.strip()
                    if para_text:
                        _add_styled_paragraph(doc, para_text, size=11, space_after=80)
    else:
        # Fall back to full_text
        full_text = policy.get("full_text", "")
        for para_text in full_text.split("\n\n"):
            para_text = para_text.strip()
            if para_text:
                _add_styled_paragraph(doc, para_text, size=11, space_after=80)

    # Drafting notes
    if policy.get("drafting_notes"):
        doc.add_paragraph()
        _add_horizontal_rule(doc)
        _add_styled_paragraph(doc, "DRAFTING NOTES", bold=True, size=10, color=COLOR_GRAY)
        _add_styled_paragraph(doc, policy["drafting_notes"], size=10, color=COLOR_GRAY, italic=True)

    # Sources used -- this draft ran the same KB retrieval + live research as
    # gap analysis, so it gets the same verification section, not just a
    # generic disclaimer.
    doc.add_paragraph()
    _add_sources_used_section(
        doc,
        policy.get("kb_sources_used"),
        policy.get("live_research_used", False),
        policy.get("verification_overall"),
    )

    # AI disclaimer
    doc.add_paragraph()
    _add_disclaimer_box(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_draft_policy_export(policy: dict) -> tuple[bytes, str]:
    """Generate a drafted policy .docx and return (bytes, filename)."""
    title = policy.get("policy_title", "Drafted Policy")
    safe_name = title.replace(" ", "_").replace("/", "-").replace("(", "").replace(")", "")[:60]
    date_str = datetime.now().strftime("%Y-%m-%d")
    file_bytes = generate_draft_policy_docx(policy)
    filename = f"{safe_name}_{date_str}.docx"
    return file_bytes, filename


# ── Clean Updated Policy Export ──

def generate_updated_policy_docx(rewritten: RewrittenPolicy) -> bytes:
    """
    Generate a clean .docx containing ONLY the rewritten/updated policy.
    Designed to look like a real policy document — title, sections, body text —
    suitable for dropping straight into a handbook or policy library.
    """
    doc = Document()

    # Standard policy-document margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.2)
        section.right_margin = Cm(2.5)

    # ── Title ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_p.add_run(rewritten.policy_title or "Updated Policy")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = COLOR_DARK_NAVY
    run.font.name = FONT_FAMILY
    title_p.paragraph_format.space_after = Pt(4)

    _add_horizontal_rule(doc)

    # ── Metadata row ──
    meta_parts = []
    if rewritten.effective_date:
        meta_parts.append(f"Effective: {rewritten.effective_date}")
    else:
        meta_parts.append("Effective: Upon adoption")
    meta_parts.append(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    if rewritten.version_note:
        meta_parts.append(rewritten.version_note)
    _add_styled_paragraph(doc, "  |  ".join(meta_parts), size=10, color=COLOR_GRAY)

    doc.add_paragraph()
    _add_horizontal_rule(doc)
    doc.add_paragraph()

    # ── Body: section-by-section ──
    sections = rewritten.sections or []
    wrote_any_body = False
    if sections:
        for sec in sections:
            section_title = (sec.section_title or "").strip()
            section_body = (sec.rewritten_text or "").strip()
            if section_title:
                _add_styled_paragraph(
                    doc, section_title.upper(), bold=True, size=12,
                    color=COLOR_DARK_NAVY, space_before=120, space_after=60,
                )
            if section_body:
                for para_text in section_body.split("\n\n"):
                    para_text = para_text.strip()
                    if para_text:
                        _add_styled_paragraph(doc, para_text, size=11, space_after=80)
                        wrote_any_body = True

    # Fallback: if no sections OR sections produced no body text, dump full_text
    if not wrote_any_body:
        for para_text in (rewritten.full_text or "").split("\n\n"):
            para_text = para_text.strip()
            if para_text:
                _add_styled_paragraph(doc, para_text, size=11, space_after=80)

    # ── AI disclaimer (kept short — this is a finished policy, not a report) ──
    doc.add_paragraph()
    _add_disclaimer_box(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_updated_policy_export(
    rewritten: "RewrittenPolicy",
    source_file_name: Optional[str] = None,
) -> tuple[bytes, str]:
    """Generate a clean updated-policy .docx and return (bytes, filename)."""
    title = rewritten.policy_title or "Updated_Policy"
    safe_name = title.replace(" ", "_").replace("/", "-").replace("(", "").replace(")", "")[:60]
    date_str = datetime.now().strftime("%Y-%m-%d")
    file_bytes = generate_updated_policy_docx(rewritten)
    filename = f"{safe_name}_Updated_{date_str}.docx"
    return file_bytes, filename


# ── Compliance Certificate ──

def _build_certificate_content(
    doc: Document,
    *,
    policy_type: str,
    regulations: List[str],
    date_str: str,
):
    """
    Builds the one-page certificate-style summary: policy type and regulations
    reviewed. Shared by the standalone certificate export and the gap analysis
    report, which now opens with this as page 1 instead of requiring a
    separate download for the same summary.
    """
    # ── Header bar ──
    hdr = doc.add_paragraph()
    hdr.paragraph_format.space_before = Pt(0)
    hdr.paragraph_format.space_after = Pt(0)
    hdr_run = hdr.add_run("  COMPLIANCE ASSESSMENT CERTIFICATE  ")
    hdr_run.bold = True
    hdr_run.font.size = Pt(13)
    hdr_run.font.color.rgb = COLOR_WHITE
    hdr.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _shade_paragraph(hdr, "1A1A2E")

    doc.add_paragraph()

    # ── Tool / issuer ──
    _add_styled_paragraph(
        doc,
        "POLICY GAP ANALYZER",
        bold=True,
        size=18,
        color=COLOR_DARK_NAVY,
        space_before=0,
        space_after=0,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add_styled_paragraph(
        doc,
        "Multi-Industry Compliance Intelligence System",
        size=10,
        color=COLOR_GRAY,
        italic=True,
        space_before=0,
        space_after=100,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    _add_horizontal_rule(doc)

    # ── Certificate title ──
    _add_styled_paragraph(
        doc,
        "CERTIFICATE OF COMPLIANCE ASSESSMENT",
        bold=True,
        size=14,
        color=COLOR_DARK_NAVY,
        space_before=80,
        space_after=20,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add_styled_paragraph(
        doc,
        f"Assessment Date: {date_str}",
        size=10,
        color=COLOR_GRAY,
        space_before=0,
        space_after=120,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # ── Policy type ──
    _add_styled_paragraph(
        doc,
        "POLICY / PROGRAM ASSESSED",
        bold=True,
        size=9,
        color=COLOR_GRAY,
        space_before=0,
        space_after=20,
    )
    _add_styled_paragraph(
        doc,
        policy_type,
        bold=True,
        size=13,
        color=COLOR_DARK_NAVY,
        space_before=0,
        space_after=140,
    )

    # ── Regulations reviewed ──
    if regulations:
        _add_styled_paragraph(
            doc,
            f"REGULATIONS & STANDARDS REVIEWED ({len(regulations)})",
            bold=True,
            size=9,
            color=COLOR_GRAY,
            space_before=60,
            space_after=20,
        )
        regs_text = " • ".join(regulations[:12])
        if len(regulations) > 12:
            regs_text += f" + {len(regulations) - 12} more"
        _add_styled_paragraph(doc, regs_text, size=9, color=COLOR_BLACK, space_after=80)

    _add_horizontal_rule(doc)

    # ── Disclaimer ──
    disclaimer = (
        "This certificate documents that the above policy or program was reviewed on the date shown "
        "using AI-assisted regulatory analysis. This assessment is informational only and does not "
        "constitute legal advice or a guarantee of regulatory compliance. All findings should be "
        "independently reviewed by qualified compliance counsel before formal adoption or regulatory filing."
    )
    _add_styled_paragraph(
        doc,
        disclaimer,
        size=8,
        color=COLOR_LIGHT_GRAY,
        italic=True,
        space_before=40,
        space_after=0,
    )


def _shade_paragraph(para, hex_color: str):
    """Apply background shading to an entire paragraph (simulated highlight box)."""
    pPr = para._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
    pPr.append(shd)


